"""Pittsfield Township Property Tax Assessment Analyzer — Streamlit App."""

import io
import os

import streamlit as st
import pandas as pd
from docx import Document
from fpdf import FPDF

from data_loader import load_all_data
from rc_parser import parse_rc_pdf, PropertyData, AssessmentYear
from analysis_engine import (
    get_ecf_trend, get_ecf_properties, get_subdivision_name,
    get_comparable_sales, compute_sales_stats, get_land_value_trends,
    check_sales_coverage, compute_ecf_adjusted_value, get_overvaluation_pct,
    generate_appeal_summary, _compute_recommended_values,
)
from charts import (
    ecf_trend_chart, ecf_distribution_chart, sales_scatter_chart,
    sales_histogram, land_trend_chart, sales_coverage_chart,
    assessment_comparison_chart, assessment_history_chart,
)

st.set_page_config(
    page_title="Pittsfield Township Tax Analyzer",
    page_icon="🏠",
    layout="wide",
)


def check_password() -> bool:
    """Show a password gate and return True if the user has authenticated."""
    if st.session_state.get("authenticated"):
        return True

    # Center a narrow column for the login form
    _, col, _ = st.columns([1, 1, 1])
    with col:
        st.markdown(S3_FOOTER_HTML, unsafe_allow_html=True)
        st.title("Pittsfield Township")
        st.caption("Property Tax Analyzer")
        with st.form("login_form"):
            password = st.text_input("Password", type="password", key="password_input")
            submitted = st.form_submit_button("Submit", type="primary", use_container_width=True)
            if submitted:
                expected = os.environ.get("APP_PASSWORD") or st.secrets.get("APP_PASSWORD", "pittsfield2026")
                if password == expected:
                    st.session_state["authenticated"] = True
                    st.rerun()
                else:
                    st.error("Incorrect password.")
    return False


def _generate_docx(text: str) -> bytes:
    """Generate a DOCX file from the appeal summary text."""
    doc = Document()
    doc.add_heading("Property Tax Appeal Petition", level=1)
    for line in text.split("\n"):
        doc.add_paragraph(line)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _generate_pdf(text: str) -> bytes:
    """Generate a PDF file from the appeal summary text."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Helvetica", "B", 16)
    pdf.cell(0, 10, "Property Tax Appeal Petition", new_x="LMARGIN", new_y="NEXT")
    pdf.ln(4)
    pdf.set_font("Courier", size=9)
    # Max chars that fit in the printable width with Courier 9pt
    usable_w = pdf.w - pdf.l_margin - pdf.r_margin
    max_chars = int(usable_w / pdf.get_string_width("W")) if pdf.get_string_width("W") > 0 else 80
    for line in text.split("\n"):
        # Truncate long separator lines (e.g. "=" * 70) to fit page width
        stripped = line.strip()
        if stripped and len(set(stripped)) == 1 and len(stripped) > max_chars:
            line = stripped[0] * max_chars
        pdf.multi_cell(0, 4, line)
    return bytes(pdf.output())


def main():
    if not check_password():
        return

    data = load_all_data()

    # --- Sidebar ---
    with st.sidebar:
        st.markdown(S3_FOOTER_HTML, unsafe_allow_html=True)
        st.title("Pittsfield Township")
        st.subheader("Property Tax Analyzer")
        st.caption("Based on official township assessment data (2024-2026)")

        st.divider()

        # RC.pdf Upload
        uploaded_file = st.file_uploader(
            "Upload your Record Card (RC.pdf)",
            type=["pdf"],
            help="Download from [BSA Online](https://bsaonline.com/?uid=193): "
                 "search your address → click 'Record Card' → save the PDF",
        )

        prop: PropertyData | None = None
        if uploaded_file is not None:
            with st.spinner("Reading your Record Card (OCR)..."):
                prop = parse_rc_pdf(uploaded_file.read())

            if prop.area_code and prop.area_code in data["all_areas"]:
                st.success(f"Parsed: **{prop.address}** ({prop.area_code})")
            elif prop.area_code:
                st.warning(f"Parsed area code **{prop.area_code}** not found in township data. "
                           "Please select manually below.")
                prop = None
            else:
                st.error("Could not parse area code from PDF. Please enter manually below.")
                prop = None

        st.divider()

        # Area code: auto-filled from PDF or manual selection
        if prop:
            area_idx = data["all_areas"].index(prop.area_code) if prop.area_code in data["all_areas"] else None
            area_code = st.selectbox(
                "ECF Area Code",
                options=data["all_areas"],
                index=area_idx,
                help="Auto-detected from your Record Card",
            )
        else:
            area_code = st.selectbox(
                "Select your ECF Area Code",
                options=data["all_areas"],
                index=None,
                placeholder="Choose an area...",
                help="Find this on your assessment notice or at bsaonline.com",
            )

        # SEV: auto-filled from PDF or manual entry
        default_sev = prop.sev_2026 if prop else 0
        sev = st.number_input(
            "Your 2026 SEV (Assessed Value)",
            min_value=0, max_value=2_000_000, value=default_sev, step=1000,
            format="%d",
            help="State Equalized Value from your 2026 assessment notice",
        )

        if sev > 0:
            st.metric("Implied TCV", f"${sev * 2:,.0f}", help="True Cash Value = SEV × 2")

        # Show parsed property details OR manual entry form
        if prop:
            st.divider()
            st.markdown("**Property Details (from PDF)**")
            details = {
                "Parcel": prop.parcel_number,
                "Style": prop.style,
                "Year Built": prop.year_built,
                "Floor Area": f"{prop.floor_area:,} SF" if prop.floor_area else "N/A",
                "Basement": f"{prop.basement_sf:,} SF" if prop.basement_sf else "N/A",
                "Condition": prop.condition,
                "Land Value": f"${prop.land_value:,}" if prop.land_value else "N/A",
            }
            for k, v in details.items():
                if v and v != "N/A" and v != "0" and v != "0 SF":
                    st.caption(f"{k}: **{v}**")
        else:
            st.divider()
            with st.expander("Property Details (manual entry)", expanded=False):
                st.caption("Optional — fills in the petition with your property info")
                m_address = st.text_input("Property Address", placeholder="e.g. 4806 PAULINA DR")
                m_parcel = st.text_input("Parcel Number", placeholder="e.g. L-12-13-311-061")
                m_style = st.selectbox("Style", options=["", "TWO-STORY", "ONE-STORY", "RANCH",
                                       "COLONIAL", "CAPE-COD", "SPLIT-LEVEL", "TRI-LEVEL",
                                       "BUNGALOW", "BI-LEVEL"], index=0)
                m_year_built = st.number_input("Year Built", min_value=0, max_value=2026,
                                                value=0, step=1, format="%d")
                m_floor_area = st.number_input("Floor Area (SF)", min_value=0, max_value=20000,
                                                value=0, step=100, format="%d")
                m_basement = st.number_input("Basement (SF)", min_value=0, max_value=10000,
                                              value=0, step=100, format="%d")
                m_condition = st.selectbox("Condition", options=["", "Excellent", "Very Good",
                                           "Good", "Average", "Fair", "Poor"], index=0)
                m_land_value = st.number_input("Land Value", min_value=0, max_value=500000,
                                                value=0, step=1000, format="%d")

                st.markdown("---")
                st.caption("**Assessment History** — from your assessment notice")
                m_taxable = st.number_input("2026 Taxable Value", min_value=0,
                                             max_value=2_000_000, value=0, step=1000,
                                             format="%d")
                # Prior year values for history chart/table
                m_sev_2025 = st.number_input("2025 SEV", min_value=0, max_value=2_000_000,
                                              value=0, step=1000, format="%d")
                m_sev_2024 = st.number_input("2024 SEV", min_value=0, max_value=2_000_000,
                                              value=0, step=1000, format="%d")
                m_sev_2023 = st.number_input("2023 SEV", min_value=0, max_value=2_000_000,
                                              value=0, step=1000, format="%d")

                # Build a PropertyData from manual inputs if address is provided
                if m_address.strip():
                    # Build assessment history from manual entries
                    m_history: list[AssessmentYear] = []
                    if m_sev_2023 > 0:
                        m_history.append(AssessmentYear(year=2023, assessed_value=m_sev_2023))
                    if m_sev_2024 > 0:
                        m_history.append(AssessmentYear(year=2024, assessed_value=m_sev_2024))
                    if m_sev_2025 > 0:
                        m_history.append(AssessmentYear(year=2025, assessed_value=m_sev_2025))
                    if sev > 0:
                        m_history.append(AssessmentYear(
                            year=2026, assessed_value=sev,
                            taxable_value=m_taxable if m_taxable > 0 else None,
                            land_value=m_land_value if m_land_value > 0 else None,
                            building_value=(sev - m_land_value) if m_land_value > 0 and sev > m_land_value else None,
                        ))

                    prop = PropertyData(
                        address=m_address.strip().upper(),
                        parcel_number=m_parcel.strip(),
                        area_code=area_code or "",
                        style=m_style,
                        year_built=m_year_built if m_year_built > 0 else 0,
                        floor_area=m_floor_area,
                        basement_sf=m_basement,
                        condition=m_condition,
                        land_value=m_land_value,
                        sev_2026=sev,
                        tcv_2026=sev * 2,
                        assessment_history=m_history,
                    )

        st.divider()
        st.markdown(
            "**Appeal Deadline:** March 10, 2026 at 5:00 PM\n\n"
            "Pittsfield Township\n"
            "6201 W. Michigan Ave\n"
            "Ann Arbor, MI 48108\n\n"
            "Phone: 734-822-3115"
        )

    # --- Main Content ---
    if not area_code:
        st.title("Pittsfield Township Property Tax Assessment Analyzer")
        st.markdown("""
        ### How to use this tool

        1. **Upload your Record Card PDF** from [BSA Online](https://bsaonline.com/?uid=193)
           - Search your address → click "Record Card" → save the PDF
           - The app will auto-detect your ECF area code and SEV
        2. **Or manually select your ECF Area Code** and enter your 2026 SEV
        3. View your personalized analysis across multiple tabs

        ### What this tool analyzes

        Using **three years of official Pittsfield Township assessment data** (2024-2026),
        this tool compares your assessment against:

        - **ECF (Economic Condition Factor)** — the township's own metric showing whether
          the cost approach over- or under-values properties in your area
        - **Comparable arm's-length sales** in your area
        - **Land value trends** and adjustment factors
        - **Sales study coverage** — whether your area is included in the assessment study

        ### Data Source

        All data comes from official Pittsfield Township documents available at
        [pittsfield-mi.gov](https://pittsfield-mi.gov/2230/Property-Assessment-Data):
        - Residential Sales Analysis (2024, 2025, 2026)
        - Residential ECF Analysis (2024, 2025, 2026)
        - Residential Land Analysis (2024, 2025, 2026)
        """)
        return

    if sev <= 0:
        st.info("Enter your 2026 SEV in the sidebar to see your personalized analysis.")
        return

    user_tcv = sev * 2
    subdivision = get_subdivision_name(data, area_code)

    # Compute all analysis
    ecf_trend = get_ecf_trend(data, area_code)
    ecf_props = get_ecf_properties(data, area_code)
    sales_df = get_comparable_sales(data, area_code)
    sales_stats = compute_sales_stats(sales_df, user_tcv)
    land_trends = get_land_value_trends(data, area_code)
    coverage = check_sales_coverage(data, area_code)

    ecf_2026 = ecf_trend.get(2026)
    ecf_adjusted = compute_ecf_adjusted_value(ecf_2026, user_tcv)
    overval_pct = get_overvaluation_pct(ecf_2026)

    # --- Header ---
    header = f"{area_code} — {subdivision}"
    if prop:
        header = f"{prop.address} | {area_code} — {subdivision}"
    st.title(header)

    # Compute recommended (moderate) values
    rec_low, rec_high = _compute_recommended_values(ecf_2026, user_tcv, sales_stats)
    rec_sev = round(rec_low / 2 / 5000) * 5000
    rec_tcv = rec_sev * 2

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("Your SEV", f"${sev:,.0f}")
    col2.metric("Implied TCV", f"${user_tcv:,.0f}")
    if ecf_2026 is not None:
        delta_str = f"-{overval_pct:.1f}% overvaluation" if overval_pct else None
        col3.metric("2026 ECF", f"{ecf_2026:.3f}", delta=delta_str, delta_color="inverse")
    else:
        col3.metric("2026 ECF", "N/A")
    no_appeal = rec_sev >= sev
    if no_appeal:
        col4.metric("Appeal", "Not Recommended",
                     delta="Assessment at/below market", delta_color="off")
    else:
        col4.metric("Recommended TCV", f"${rec_tcv:,.0f}",
                     delta=f"${rec_tcv - user_tcv:,.0f}", delta_color="normal")

    # Key alert banner
    if no_appeal:
        st.success(
            f"Your assessment appears to be at or below market value. "
            f"Median of comparable sales: **\\${sales_stats['median']:,.0f}** "
            f"vs your TCV: **\\${user_tcv:,.0f}**. "
            f"An appeal is not recommended."
        )
    else:
        alerts = []
        if sales_stats.get("count", 0) > 0 and sales_stats["delta_from_median"] > 0:
            alerts.append(f"Your TCV is **\\${sales_stats['delta_from_median']:,.0f}** above the area median sale price (\\${sales_stats['median']:,.0f})")
        if overval_pct and overval_pct > 5:
            alerts.append(f"ECF of {ecf_2026:.3f} indicates cost approach overvalues by **{overval_pct:.1f}%**")
        if alerts:
            st.warning(" | ".join(alerts))

    # --- Tabs ---
    tab_names = ["ECF Analysis", "Comparable Sales", "Land Values", "Sales Coverage"]
    if prop and prop.assessment_history:
        tab_names.append("Assessment History")
    tab_names.append("Appeal Summary")
    tabs = st.tabs(tab_names)
    tab_idx = 0

    # --- Tab: ECF Analysis ---
    with tabs[tab_idx]:
        col_left, col_right = st.columns(2)
        with col_left:
            st.plotly_chart(ecf_trend_chart(ecf_trend), use_container_width=True)
        with col_right:
            st.plotly_chart(ecf_distribution_chart(ecf_props), use_container_width=True)

        # ECF Explanation
        if ecf_2026 is not None:
            if ecf_2026 < 1.0:
                st.markdown(f"""
                **What this means:** An ECF of **{ecf_2026:.3f}** means the township's cost approach
                values properties in {area_code} at **{overval_pct:.1f}% more** than what they actually
                sell for. When applied to your assessment:

                - Cost-based TCV × ECF = Market-adjusted TCV
                - **${user_tcv:,.0f} × {ecf_2026:.3f} = ${ecf_adjusted:,.0f}**

                This suggests your assessment may overstate market value by **${user_tcv - ecf_adjusted:,.0f}**.
                """)
            else:
                st.markdown(f"""
                **What this means:** An ECF of **{ecf_2026:.3f}** indicates the cost approach
                does not overvalue properties in {area_code}. The assessment appears supported
                by market evidence from the ECF perspective.
                """)

        # Individual ECF table
        if ecf_props:
            st.subheader("Individual Property ECFs in Your Area")
            ecf_df = pd.DataFrame(ecf_props)
            ecf_df = ecf_df.rename(columns={
                "year": "Year", "address": "Address", "sale_price": "Sale Price",
                "cost_man": "Cost Manual", "ecf": "ECF",
            })
            display_cols = ["Year", "Address", "Sale Price", "Cost Manual", "ECF"]
            display_df = ecf_df[[c for c in display_cols if c in ecf_df.columns]].copy()
            for col in ["Sale Price", "Cost Manual"]:
                if col in display_df.columns:
                    display_df[col] = display_df[col].apply(
                        lambda x: f"${x:,.0f}" if pd.notna(x) else "N/A"
                    )
            if "ECF" in display_df.columns:
                display_df["ECF"] = display_df["ECF"].apply(
                    lambda x: f"{x:.3f}" if pd.notna(x) else "N/A"
                )
            st.dataframe(display_df, use_container_width=True, hide_index=True)
    tab_idx += 1

    # --- Tab: Comparable Sales ---
    with tabs[tab_idx]:
        if sales_df.empty:
            st.warning(f"No arm's-length completed-home sales found for {area_code} in the 2024-2026 data. "
                       "This means the township has no comparable sales to support the assessment.")
        else:
            # Value comparison bar
            median_sale = sales_stats.get("median")
            st.plotly_chart(assessment_comparison_chart(user_tcv, ecf_adjusted, median_sale),
                            use_container_width=True)

            col_left, col_right = st.columns(2)
            with col_left:
                st.plotly_chart(sales_scatter_chart(sales_df, user_tcv), use_container_width=True)
            with col_right:
                st.plotly_chart(sales_histogram(sales_df, user_tcv), use_container_width=True)

            # Stats summary
            st.subheader("Sales Statistics")
            scol1, scol2, scol3, scol4 = st.columns(4)
            scol1.metric("Total Sales", sales_stats["count"])
            scol2.metric("Median Price", f"${sales_stats['median']:,.0f}")
            scol3.metric("Mean Price", f"${sales_stats['mean']:,.0f}")
            if sales_stats["delta_from_median"] > 0:
                scol4.metric("Your TCV vs Median", f"+${sales_stats['delta_from_median']:,.0f}",
                             delta=f"{sales_stats['pct_above_tcv']:.0f}% of sales below your TCV",
                             delta_color="inverse")
            else:
                scol4.metric("Your TCV vs Median", f"-${abs(sales_stats['delta_from_median']):,.0f}")

            # Sales data table
            st.subheader("All Comparable Sales")
            display_sales = sales_df[["Street_Address", "Sale_Date", "Sale_Price", "Adj_Sale", "_year"]].copy()
            display_sales = display_sales.rename(columns={
                "Street_Address": "Address", "Sale_Date": "Date",
                "Sale_Price": "Sale Price", "Adj_Sale": "Adj. Sale", "_year": "Study Year",
            })
            display_sales["vs TCV"] = display_sales["Adj. Sale"].apply(
                lambda x: f"${x - user_tcv:+,.0f}" if pd.notna(x) else "N/A"
            )
            for col in ["Sale Price", "Adj. Sale"]:
                display_sales[col] = display_sales[col].apply(
                    lambda x: f"${x:,.0f}" if pd.notna(x) else "N/A"
                )
            st.dataframe(display_sales, use_container_width=True, hide_index=True)
    tab_idx += 1

    # --- Tab: Land Values ---
    with tabs[tab_idx]:
        if not land_trends:
            st.warning(f"No land value data found for {area_code}.")
        else:
            st.plotly_chart(land_trend_chart(land_trends), use_container_width=True)

            # Land trends table
            st.subheader("Land Value Details")
            lt_rows = []
            for lt in land_trends:
                factor = lt.get("adjust_factor")
                lt_rows.append({
                    "Year": lt["year"],
                    "Prior Land Value": f"${lt['prior_lv']:,.0f}" if lt.get("prior_lv") else "N/A",
                    "Adjust Factor": f"{factor:.4f}" if factor else "N/A",
                    "Change": f"{(factor - 1) * 100:+.1f}%" if factor else "N/A",
                    "Current Land Value": f"${lt['current_lv']:,.0f}" if lt.get("current_lv") else "N/A",
                })
            st.dataframe(pd.DataFrame(lt_rows), use_container_width=True, hide_index=True)

            # Cumulative change
            first_prior = next((lt["prior_lv"] for lt in land_trends if lt.get("prior_lv")), None)
            last_current = next((lt["current_lv"] for lt in reversed(land_trends) if lt.get("current_lv")), None)
            if first_prior and last_current and first_prior > 0:
                total_pct = (last_current - first_prior) / first_prior * 100
                st.info(f"Total land value change over 3 years: **${first_prior:,.0f} -> ${last_current:,.0f}** "
                        f"(**{total_pct:+.1f}%**)")
    tab_idx += 1

    # --- Tab: Sales Coverage ---
    with tabs[tab_idx]:
        st.plotly_chart(sales_coverage_chart(coverage), use_container_width=True)

        cov_2026 = coverage.get(2026, 0)
        cov_2025 = coverage.get(2025, 0)
        cov_2024 = coverage.get(2024, 0)

        if cov_2026 == 0 and (cov_2024 > 0 or cov_2025 > 0):
            st.error(f"""
            **CRITICAL: {area_code} was dropped from the 2026 sales study.**

            - 2024 study: {cov_2024} sales from your area
            - 2025 study: {cov_2025} sales from your area
            - 2026 study: **0 sales** (area dropped)

            The township increased assessments for {area_code} in 2026 without any supporting
            market sales evidence. This is a strong argument for appeal.
            """)
        elif cov_2026 == 0:
            st.warning(f"No sales from {area_code} appear in any year's sales study. "
                       "The assessment relies entirely on the cost approach with no market validation.")
        else:
            st.success(f"{area_code} has {cov_2026} sales in the 2026 study. "
                       "The township has market evidence for the assessment.")
    tab_idx += 1

    # --- Tab: Assessment History (only if RC.pdf uploaded) ---
    if prop and prop.assessment_history:
        with tabs[tab_idx]:
            st.plotly_chart(assessment_history_chart(prop.assessment_history),
                            use_container_width=True)

            # Year-over-year changes
            st.subheader("Year-over-Year Assessment Changes")
            hist_rows = []
            prev_assessed = None
            for h in prop.assessment_history:
                row = {
                    "Year": h.year,
                    "Land Value": f"${h.land_value:,}" if h.land_value else "N/A",
                    "Building Value": f"${h.building_value:,}" if h.building_value else "N/A",
                    "Assessed (SEV)": f"${h.assessed_value:,}" if h.assessed_value else "N/A",
                    "Taxable": f"${h.taxable_value:,}" if h.taxable_value else "N/A",
                }
                if prev_assessed and h.assessed_value:
                    chg = (h.assessed_value - prev_assessed) / prev_assessed * 100
                    row["SEV Change"] = f"{chg:+.1f}%"
                else:
                    row["SEV Change"] = ""
                prev_assessed = h.assessed_value
                hist_rows.append(row)
            st.dataframe(pd.DataFrame(hist_rows), use_container_width=True, hide_index=True)

            # Cost approach details from RC.pdf
            if prop.total_base_new or prop.total_depr_cost:
                st.subheader("Cost Approach (from Record Card)")
                ccol1, ccol2, ccol3 = st.columns(3)
                if prop.total_base_new:
                    ccol1.metric("Total Base New (Cost)", f"${prop.total_base_new:,}")
                if prop.total_depr_cost:
                    ccol2.metric("Total Depr. Cost", f"${prop.total_depr_cost:,}")
                if prop.estimated_tcv_cost:
                    ccol3.metric("Est. TCV (Cost Approach)", f"${prop.estimated_tcv_cost:,}")

                if prop.ecf and prop.estimated_tcv_cost:
                    st.markdown(f"""
                    The assessor's cost approach produces an estimated TCV of **${prop.estimated_tcv_cost:,}**
                    after applying the ECF of **{prop.ecf:.3f}**. Compare this to the official 2026 Est TCV
                    of **${prop.tcv_2026:,}** (SEV × 2 = ${sev * 2:,}).
                    """)
        tab_idx += 1

    # --- Tab: Appeal Summary ---
    with tabs[tab_idx]:
        summary = generate_appeal_summary(
            area_code, subdivision, sev,
            ecf_trend, sales_stats, land_trends, coverage, ecf_props,
            sales_df=sales_df, prop=prop,
        )

        if no_appeal:
            st.error("**Appeal Not Recommended**")
            st.markdown(
                f"The median of comparable sales (**\\${sales_stats['median']:,.0f}**) "
                f"is at or above your current TCV (**\\${user_tcv:,.0f}**). "
                f"Filing an appeal is unlikely to result in a reduction and could "
                f"risk the Board increasing your assessment."
            )
            st.divider()
            st.caption("Analysis details:")
            st.code(summary, language=None)
        else:
            st.code(summary, language=None)

            base_name = f"L4035_Petition_{area_code}_{sev}"
            if prop and prop.address:
                addr_slug = prop.address.replace(" ", "_").replace(",", "")
                base_name = f"L4035_Petition_{addr_slug}"

            dl1, dl2, dl3 = st.columns(3)
            with dl1:
                st.download_button(
                    label="Download as Text",
                    data=summary,
                    file_name=f"{base_name}.txt",
                    mime="text/plain",
                    use_container_width=True,
                )
            with dl2:
                st.download_button(
                    label="Download as DOCX",
                    data=_generate_docx(summary),
                    file_name=f"{base_name}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True,
                )
            with dl3:
                st.download_button(
                    label="Download as PDF",
                    data=_generate_pdf(summary),
                    file_name=f"{base_name}.pdf",
                    mime="application/pdf",
                    use_container_width=True,
                )

            st.divider()
            st.subheader("Next Steps")
            st.markdown("""
            1. **Pull your property record card** from [BSA Online](https://bsaonline.com/?uid=193)
               — verify all physical characteristics are correct
            2. **Request the assessor's sales study** used to value your property
            3. **Complete Form L-4035** (Petition to Board of Review) and submit with this analysis
            4. **Submit by March 10, 2026 at 5:00 PM** to:
               - Pittsfield Township, 6201 W. Michigan Ave, Ann Arbor, MI 48108
               - Or call 734-822-3115 to schedule a hearing
            5. **If denied:** File with Michigan Tax Tribunal by July 31, 2026
            """)


S3_FOOTER_HTML = """
<style>
@keyframes gradient-shift {
  0%, 100% { background-position: 0% 50%; }
  50% { background-position: 100% 50%; }
}
@keyframes shimmer {
  0% { background-position: -200% center; }
  100% { background-position: 200% center; }
}
@keyframes sparkle-1 {
  0%, 100% { opacity: 0; transform: scale(0) rotate(0deg); }
  50% { opacity: 1; transform: scale(1) rotate(180deg); }
}
@keyframes sparkle-2 {
  0%, 100% { opacity: 0; transform: scale(0) rotate(0deg); }
  33% { opacity: 1; transform: scale(1) rotate(120deg); }
}
@keyframes sparkle-3 {
  0%, 100% { opacity: 0; transform: scale(0) rotate(0deg); }
  66% { opacity: 1; transform: scale(1) rotate(240deg); }
}
.s3-footer {
  text-align: center;
  padding: 1rem 0 0.5rem 0;
  font-size: 0.8rem;
  color: #888;
}
.s3-footer a {
  text-decoration: none;
  position: relative;
  display: inline-block;
}
.s3-footer .s3-glow {
  position: absolute;
  inset: 0;
  background: linear-gradient(90deg, rgba(124,58,237,0.15), rgba(219,39,119,0.15), rgba(6,182,212,0.15));
  border-radius: 12px;
  filter: blur(12px);
  transition: filter 0.3s;
}
.s3-footer a:hover .s3-glow {
  filter: blur(20px);
}
.s3-footer .s3-glass {
  position: absolute;
  inset: 0;
  background: rgba(255,255,255,0.08);
  backdrop-filter: blur(4px);
  border-radius: 12px;
  border: 1px solid rgba(255,255,255,0.15);
  box-shadow: 0 4px 12px rgba(0,0,0,0.1);
}
.s3-footer .s3-shimmer {
  position: absolute;
  inset: 0;
  border-radius: 12px;
  background: linear-gradient(105deg, transparent 40%, rgba(255,255,255,0.25) 50%, transparent 60%);
  background-size: 200% 100%;
  animation: shimmer 3s linear infinite;
}
.s3-footer .s3-text {
  position: relative;
  display: inline-block;
  padding: 2px 14px;
  font-weight: 700;
  font-size: 0.85rem;
  background: linear-gradient(90deg, #7c3aed, #db2777, #06b6d4, #7c3aed);
  background-size: 300% 100%;
  -webkit-background-clip: text;
  -webkit-text-fill-color: transparent;
  background-clip: text;
  animation: gradient-shift 3s ease infinite;
}
.s3-footer .sparkle {
  position: absolute;
  font-size: 0.6rem;
}
.s3-footer .sparkle-1 {
  top: -4px; right: -4px;
  color: #facc15;
  animation: sparkle-1 2s ease-in-out infinite;
}
.s3-footer .sparkle-2 {
  bottom: -4px; left: -4px;
  color: #60a5fa;
  animation: sparkle-2 2s ease-in-out infinite 0.3s;
}
.s3-footer .sparkle-3 {
  top: -2px; left: 50%;
  color: #f472b6;
  animation: sparkle-3 2s ease-in-out infinite 0.6s;
}
</style>
<div class="s3-footer">
  <span>Designed by</span>
  <a href="https://www.s3techinnovations.com/" target="_blank" rel="noopener noreferrer">
    <span class="s3-glow"></span>
    <span class="s3-glass"></span>
    <span class="s3-shimmer"></span>
    <span class="s3-text">S3 Tech Innovations</span>
    <span class="sparkle sparkle-1">&#10022;</span>
    <span class="sparkle sparkle-2">&#10022;</span>
    <span class="sparkle sparkle-3">&#10022;</span>
  </a>
</div>
"""


if __name__ == "__main__":
    main()
